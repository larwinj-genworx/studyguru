from __future__ import annotations

from pathlib import Path

from src.schemas.study_material import ConceptContentPack

class DocxRenderer:
    def render(
        self,
        output_dir: Path,
        subject_name: str,
        grade_level: str,
        concept_packs: list[ConceptContentPack],
    ) -> Path:
        from docx import Document

        document = Document()
        document.add_heading(f"{subject_name} Study Notes", level=0)
        document.add_paragraph(f"Grade/Level: {grade_level}")

        for pack in concept_packs:
            document.add_heading(pack.concept_name, level=1)
            document.add_paragraph(pack.definition)
            document.add_paragraph(pack.intuition)

            if pack.stepwise_breakdown_required and pack.key_steps:
                document.add_heading("Key Steps", level=2)
                for step in pack.key_steps:
                    document.add_paragraph(step, style="List Bullet")

            if pack.examples:
                document.add_heading("Worked Examples", level=2)
                for example in pack.examples:
                    document.add_paragraph(example, style="List Number")

            if pack.common_mistakes:
                document.add_heading("Common Mistakes", level=2)
                for mistake in pack.common_mistakes:
                    document.add_paragraph(mistake, style="List Bullet")

            if pack.recap:
                document.add_heading("Quick Revision", level=2)
                for line in pack.recap:
                    document.add_paragraph(line, style="List Bullet")

        output_path = output_dir / "study_material.docx"
        document.save(str(output_path))
        return output_path
